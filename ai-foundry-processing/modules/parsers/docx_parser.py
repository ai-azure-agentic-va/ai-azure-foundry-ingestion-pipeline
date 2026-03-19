"""DOCX parser using python-docx - no Azure Document Intelligence."""

import io
import logging
from docx import Document
from .base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def parse(self, file_bytes: bytes) -> ParseResult:
        logger.info(f"[DocxParser] Parsing DOCX ({len(file_bytes)} bytes)")
        doc = Document(io.BytesIO(file_bytes))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        tables_text = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                tables_text.append("\n".join(" | ".join(row) for row in rows))

        full_text = "\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n" + "\n\n".join(tables_text)

        logger.info(f"[DocxParser] Extracted {len(paragraphs)} paragraphs, {len(tables_text)} tables, {len(full_text)} chars")
        return ParseResult(
            full_text=full_text,
            pages=[],
            page_count=1,
            metadata={"format": "docx"},
        )
